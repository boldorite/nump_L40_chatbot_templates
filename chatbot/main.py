import os
import json
import re
import uuid
import subprocess
from pathlib import Path
from datetime import datetime

import httpx
from dotenv import load_dotenv
from fastapi import FastAPI, Request
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse, HTMLResponse
from jinja2 import Template

load_dotenv()

app = FastAPI(title="NUMP Chatbot with Templates", version="1.0.0")

OLLAMA_BASE = os.getenv("LLM_BASE_URL", "http://localhost:11434").rstrip("/v1").rstrip("/")
MODEL = os.getenv("MODEL", "qwen3.5:27b-q8_0")
DATA_DIR = Path(os.getenv("DATA_DIR", "../data"))
EXPORT_DIR = DATA_DIR / "exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

sessions: dict[str, dict] = {}


# ── LLM ──
def call_llm_stream(messages: list[dict]):
    payload = {
        "model": MODEL, "messages": messages, "stream": True, "think": False,
        "options": {"num_predict": 4096, "temperature": 0.7},
    }
    with httpx.stream("POST", f"{OLLAMA_BASE}/api/chat", json=payload, timeout=180.0) as resp:
        for line in resp.iter_lines():
            if line:
                try:
                    chunk = json.loads(line)
                    content = chunk.get("message", {}).get("content", "")
                    if content:
                        yield content
                    if chunk.get("done"):
                        break
                except json.JSONDecodeError:
                    continue


def call_llm(messages: list[dict], max_tokens: int = 4096) -> str:
    payload = {
        "model": MODEL, "messages": messages, "stream": False, "think": False,
        "options": {"num_predict": max_tokens, "temperature": 0.7},
    }
    resp = httpx.post(f"{OLLAMA_BASE}/api/chat", json=payload, timeout=180.0)
    resp.raise_for_status()
    return resp.json().get("message", {}).get("content", "")


# ── 세션 ──
def get_or_create_session(sid: str | None) -> tuple[str, dict]:
    if sid and sid in sessions:
        return sid, sessions[sid]
    new_sid = sid or str(uuid.uuid4())[:8]
    sessions[new_sid] = {
        "messages": [],
        "folder_name": None,       # 선택된 폴더명
        "template_on": False,
        "active_schema": None,     # 클라이언트에서 업로드된 schema
        "active_template_j2": None,# 클라이언트에서 업로드된 template.j2
        "generated_data": {},
    }
    return new_sid, sessions[new_sid]


# ── 프롬프트 ──
SYSTEM_NORMAL = """당신은 친절하고 유능한 한국어 AI 어시스턴트 'NUMP 챗봇'입니다.
사용자의 질문에 정확하고 간결하게 답변하세요. 마크다운 형식을 사용할 수 있습니다.
수학 수식은 반드시 LaTeX 형식으로 $...$ (인라인) 또는 $$...$$ (블록)으로 감싸서 작성하세요.
예: 근의 공식은 $$x = \\frac{-b \\pm \\sqrt{b^2 - 4ac}}{2a}$$ 이렇게 작성."""


def build_template_prompt(schema: dict) -> str:
    parts = []
    for sec in schema.get("sections", []):
        parts.append(f"\n[{sec['section_name']}]")
        for f in sec.get("fields", []):
            req = "필수" if f.get("required") else "선택"
            extra = ""
            if f.get("type") == "table" and f.get("table_columns"):
                extra = f" (테이블: {', '.join(f['table_columns'])})"
            parts.append(f"  - {f['label']} (id:{f['field_id']}, {f.get('type','text')}, {req}){extra}")
    return f"""당신은 친절하고 유능한 한국어 AI 어시스턴트 'NUMP 챗봇'입니다.

현재 서식 모드 ON — 참조 서식: **{schema.get('form_name','')}**
설명: {schema.get('description','')}

서식 포맷 구조:
{''.join(parts)}

## 역할
1. 사용자와 자연스럽게 대화하세요.
2. 이 서식 포맷을 참고하여 **새 문서**를 만드는 것이 목표입니다.
3. 대화 중 사용자가 제공하는 정보를 기억하세요.
4. 일반 질문에도 정상적으로 답변하세요.

## 절대 금지 (매우 중요)
- 서식 필드 내용을 채팅에 절대 출력하지 마세요.
- "소속: XX, 성명: XX" 같은 필드 나열 금지.
- JSON, 표, 서식 전체/부분 출력 금지.
- 문서 내용은 오른쪽 미리보기 패널에서 자동 표시됩니다.
- 채팅에서는 1~2문장으로만 답하세요.
- 좋은 예: "알겠습니다, 반영했습니다." / "연구 기간도 알려주세요."
- 나쁜 예: 필드 목록 나열, 서식 내용 반복 출력"""


GENERATE_DOC_PROMPT = """아래 대화 내용과 서식 포맷을 참고하여 새 문서의 각 필드 값을 작성하세요.

서식명: {form_name}
필드 목록:
{fields_desc}

대화 내용:
{conversation}

## 지시사항
- 대화에서 언급된 정보를 해당 필드에 넣으세요.
- 대화에서 직접 언급되지 않았지만 맥락에서 유추 가능한 정보는 합리적으로 작성하세요.
  예: 연구과제명이 언급되었으면 연구목적/연구내용/기대효과를 과제명에 맞게 작성.
  예: 비상연락처가 언급되지 않았어도 가상의 합리적인 정보로 채우세요.
- 빈 문자열("")로 두지 마세요. 가능한 모든 필드를 채우세요.
- table 타입 필드는 리스트로: [{{"컬럼1":"값","컬럼2":"값"}}, ...] (최소 1행 이상)
- 정보가 부족해도 반드시 모든 필드를 포함하고 합리적으로 채우세요.
- 반드시 아래 JSON 형식만 반환. 다른 텍스트 없이:

{{"field_id": "값", ...}}"""


# ══════════════ API ══════════════

@app.get("/api/model")
async def get_model():
    return {"model": MODEL}

@app.post("/api/session/upload-template")
async def upload_template(request: Request):
    """클라이언트에서 읽은 schema.json + template.j2를 세션에 저장"""
    body = await request.json()
    sid = body.get("session_id")
    session_id, session = get_or_create_session(sid)

    session["active_schema"] = body.get("schema")
    session["active_template_j2"] = body.get("template_j2")
    session["folder_name"] = body.get("folder_name", "")
    session["template_on"] = True
    session["generated_data"] = {}

    form_name = ""
    if session["active_schema"]:
        form_name = session["active_schema"].get("form_name", "")

    return {"session_id": session_id, "form_name": form_name}


@app.post("/api/delete-template")
async def delete_template(request: Request):
    """서식 폴더 및 registry 항목 삭제"""
    body = await request.json()
    folder_name = body.get("folder_name")
    if not folder_name:
        return JSONResponse({"error": "folder_name 필요"}, status_code=400)

    # templates 폴더 삭제
    import shutil
    tpl_dir = DATA_DIR / "templates" / folder_name
    if tpl_dir.exists():
        shutil.rmtree(tpl_dir)

    # registry.json에서 해당 항목 제거
    registry_path = DATA_DIR / "registry.json"
    if registry_path.exists():
        registry = json.loads(registry_path.read_text(encoding="utf-8"))
        forms = registry.get("forms", {})
        to_remove = [num for num, form in forms.items()
                     if Path(form.get("schema_path", "")).parent.name == folder_name]
        for num in to_remove:
            del forms[num]
        registry_path.write_text(json.dumps(registry, ensure_ascii=False, indent=2), encoding="utf-8")

    return {"deleted": True, "folder": folder_name}


@app.post("/api/render-preview")
async def render_preview(request: Request):
    """template.j2에 데이터를 넣어 렌더링한 HTML 반환"""
    body = await request.json()
    sid = body.get("session_id")
    filled = body.get("filled_data", {})

    session = sessions.get(sid)
    if not session:
        return HTMLResponse("<p>세션 없음</p>")

    schema = session.get("active_schema")
    tpl_str = session.get("active_template_j2")
    if not tpl_str:
        return HTMLResponse("<p>템플릿 없음</p>")

    # 테이블 필드 문자열→리스트 변환
    if schema:
        for sec in schema.get("sections", []):
            for f in sec.get("fields", []):
                if f.get("type") == "table":
                    fid = f["field_id"]
                    val = filled.get(fid)
                    if isinstance(val, str):
                        try:
                            parsed = json.loads(val)
                            if isinstance(parsed, list):
                                filled[fid] = parsed
                        except Exception:
                            pass
                    elif val is None:
                        filled[fid] = []

    tpl = Template(tpl_str)
    render_data = {
        **filled,
        "form_name": schema.get("form_name", "") if schema else "",
        "created_date": datetime.now().strftime("%Y-%m-%d"),
    }
    return HTMLResponse(tpl.render(**render_data))


@app.post("/api/session/toggle-template")
async def toggle_template(request: Request):
    body = await request.json()
    sid = body.get("session_id")
    enabled = body.get("enabled", False)
    if sid and sid in sessions:
        sessions[sid]["template_on"] = enabled
        if not enabled:
            sessions[sid]["generated_data"] = {}
        return {"template_on": enabled}
    return JSONResponse({"error": "세션 없음"}, status_code=404)


@app.post("/api/session/reset")
async def reset_session(request: Request):
    body = await request.json()
    sid = body.get("session_id")
    if sid and sid in sessions:
        del sessions[sid]
    return {"reset": True}


@app.post("/api/chat")
async def chat(request: Request):
    body = await request.json()
    msg = body.get("message", "")
    sid = body.get("session_id")
    session_id, session = get_or_create_session(sid)

    # 시스템 프롬프트
    if session["template_on"] and session.get("active_schema"):
        sys_prompt = build_template_prompt(session["active_schema"])
    else:
        sys_prompt = SYSTEM_NORMAL

    session["messages"].append({"role": "user", "content": msg})
    llm_messages = [{"role": "system", "content": sys_prompt}] + session["messages"][-20:]

    async def generate():
        full = []
        try:
            for chunk in call_llm_stream(llm_messages):
                full.append(chunk)
                yield f"data: {json.dumps({'content': chunk}, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"

        session["messages"].append({"role": "assistant", "content": "".join(full)})
        meta = {
            "done": True, "session_id": session_id,
            "template_on": session["template_on"],
            "folder_name": session.get("folder_name"),
        }
        yield f"data: {json.dumps(meta, ensure_ascii=False)}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Session-Id": session_id})


@app.post("/api/generate-doc")
async def generate_doc(request: Request):
    body = await request.json()
    sid = body.get("session_id")
    session = sessions.get(sid)
    if not session:
        return JSONResponse({"error": "세션 없음"}, status_code=404)

    schema = session.get("active_schema")
    if not schema:
        return JSONResponse({"error": "서식 없음. 서식 폴더를 먼저 선택하세요."}, status_code=404)

    fields_lines = []
    for sec in schema.get("sections", []):
        for f in sec.get("fields", []):
            extra = ""
            if f.get("type") == "table" and f.get("table_columns"):
                extra = f" (테이블 컬럼: {', '.join(f['table_columns'])})"
            fields_lines.append(f"{f['field_id']}: {f['label']} ({f.get('type','text')}){extra}")

    conversation = "\n".join([
        f"{'사용자' if m['role']=='user' else 'AI'}: {m['content']}"
        for m in session["messages"][-20:]
    ])

    prompt = GENERATE_DOC_PROMPT.format(
        form_name=schema.get("form_name", ""),
        fields_desc="\n".join(fields_lines),
        conversation=conversation,
    )

    result = call_llm([{"role": "user", "content": prompt}], max_tokens=4096)

    try:
        clean = re.sub(r"```json|```", "", result).strip()
        depth = 0
        start = -1
        for i, ch in enumerate(clean):
            if ch == '{':
                if depth == 0: start = i
                depth += 1
            elif ch == '}':
                depth -= 1
                if depth == 0 and start >= 0:
                    try:
                        data = json.loads(clean[start:i+1])
                        session["generated_data"] = data
                        return {"filled_data": data}
                    except json.JSONDecodeError:
                        start = -1
        return JSONResponse({"error": "JSON 파싱 실패", "raw": result[:500]}, status_code=500)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/api/export")
async def export_doc(request: Request):
    body = await request.json()
    sid = body.get("session_id")
    fmt = body.get("format", "pdf").lower()
    filled = body.get("filled_data", {})

    session = sessions.get(sid)
    if not session:
        return JSONResponse({"error": "세션 없음"}, status_code=404)

    schema = session.get("active_schema")
    tpl_str = session.get("active_template_j2")

    if not schema or not tpl_str:
        return JSONResponse({"error": "서식/템플릿 없음"}, status_code=404)

    # 테이블 필드: 문자열이면 리스트로 변환
    for sec in schema.get("sections", []):
        for f in sec.get("fields", []):
            if f.get("type") == "table":
                fid = f["field_id"]
                val = filled.get(fid)
                if isinstance(val, str):
                    try:
                        parsed = json.loads(val)
                        if isinstance(parsed, list):
                            filled[fid] = parsed
                    except Exception:
                        pass
                elif val is None:
                    filled[fid] = []

    # Jinja2 렌더링
    tpl = Template(tpl_str)
    render_data = {
        **filled,
        "form_name": schema.get("form_name", ""),
        "created_date": datetime.now().strftime("%Y-%m-%d"),
    }
    html_content = tpl.render(**render_data)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"{schema.get('form_id', 'doc')}_{ts}"

    if fmt == "html":
        out = EXPORT_DIR / f"{base}.html"
        out.write_text(html_content, encoding="utf-8")
        return FileResponse(str(out), filename=f"{base}.html", media_type="text/html")

    elif fmt == "pdf":
        # Playwright Chromium으로 HTML→PDF (HTML 그대로 보존)
        html_tmp = EXPORT_DIR / f"{base}.html"
        html_tmp.write_text(html_content, encoding="utf-8")
        out_path = EXPORT_DIR / f"{base}.pdf"
        try:
            from playwright.async_api import async_playwright
            async with async_playwright() as p:
                browser = await p.chromium.launch()
                page = await browser.new_page()
                await page.goto(f"file://{html_tmp.resolve()}")
                await page.pdf(path=str(out_path), format="A4", margin={"top":"15mm","bottom":"15mm","left":"15mm","right":"15mm"}, print_background=True)
                await browser.close()
            html_tmp.unlink(missing_ok=True)
            return FileResponse(str(out_path), filename=f"{base}.pdf", media_type="application/pdf")
        except Exception as e:
            html_tmp.unlink(missing_ok=True)
            return JSONResponse({"error": f"PDF 변환 실패: {e}"}, status_code=500)

    elif fmt == "docx":
        out_path = EXPORT_DIR / f"{base}.docx"
        try:
            _build_docx(schema, filled, str(out_path))
            return FileResponse(str(out_path), filename=f"{base}.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            return JSONResponse({"error": f"DOCX 생성 실패: {e}"}, status_code=500)

    return JSONResponse({"error": f"지원하지 않는 형식: {fmt}"}, status_code=400)


def _build_docx(schema: dict, filled: dict, output_path: str):
    """schema + filled data로 서식 형태의 DOCX 직접 생성"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn

    doc = Document()

    # 기본 스타일 — Noto Sans CJK KR (한글 지원)
    style = doc.styles['Normal']
    style.font.name = 'Noto Sans CJK KR'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK KR')
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run(schema.get("form_name", "서식"))
    run.bold = True
    run.font.size = Pt(18)

    # 작성일/문서번호
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    run = meta.add_run(f"작성일: {filled.get('created_date', datetime.now().strftime('%Y-%m-%d'))}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for section in schema.get("sections", []):
        fields = section.get("fields", [])
        if not fields:
            continue

        # 섹션 제목
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(10)
        sp.paragraph_format.space_after = Pt(2)
        run = sp.add_run(section["section_name"])
        run.bold = True
        run.font.size = Pt(11)
        # 섹션 배경색 효과 (왼쪽 바)
        shading = sp.paragraph_format.element
        pPr = shading.get_or_add_pPr()
        pBdr = pPr.makeelement(qn('w:pBdr'), {})
        left = pBdr.makeelement(qn('w:left'), {
            qn('w:val'): 'single', qn('w:sz'): '12',
            qn('w:space'): '4', qn('w:color'): '4A5568'
        })
        pBdr.append(left)
        pPr.append(pBdr)

        # 테이블 타입 필드와 일반 필드 분리
        normal_fields = [f for f in fields if f.get("type") != "table"]
        table_fields = [f for f in fields if f.get("type") == "table"]

        # 일반 필드 → 2열 테이블 (라벨 | 값)
        if normal_fields:
            # row_group 처리
            rows_data = []
            i = 0
            while i < len(normal_fields):
                f = normal_fields[i]
                ftype = f.get("type", "text")
                rg = f.get("row_group")

                if ftype == "textarea":
                    # textarea는 전체 너비
                    rows_data.append([f])
                    i += 1
                elif rg is not None and i + 1 < len(normal_fields) and normal_fields[i+1].get("row_group") == rg:
                    # 같은 row_group → 한 행에 2필드
                    rows_data.append([f, normal_fields[i+1]])
                    i += 2
                else:
                    rows_data.append([f])
                    i += 1

            tbl = doc.add_table(rows=0, cols=4)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 열 너비 설정
            for row_fields in rows_data:
                row = tbl.add_row()
                if len(row_fields) == 1:
                    f = row_fields[0]
                    fid = f["field_id"]
                    val = str(filled.get(fid, ""))
                    ftype = f.get("type", "text")

                    if ftype == "textarea":
                        # 라벨 행
                        row.cells[0].merge(row.cells[3])
                        p = row.cells[0].paragraphs[0]
                        run = p.add_run(f["label"])
                        run.bold = True
                        run.font.size = Pt(9)
                        _shade_cell(row.cells[0], "F8F8F8")
                        # 값 행
                        val_row = tbl.add_row()
                        val_row.cells[0].merge(val_row.cells[3])
                        p = val_row.cells[0].paragraphs[0]
                        p.add_run(val).font.size = Pt(10)
                        val_row.cells[0].paragraphs[0].paragraph_format.space_after = Pt(6)
                    else:
                        # 라벨(1칸) + 값(3칸)
                        row.cells[1].merge(row.cells[3])
                        p = row.cells[0].paragraphs[0]
                        run = p.add_run(f["label"])
                        run.bold = True
                        run.font.size = Pt(9)
                        _shade_cell(row.cells[0], "F8F8F8")
                        row.cells[1].paragraphs[0].add_run(val).font.size = Pt(10)
                else:
                    # 2필드 한 행: 라벨1|값1|라벨2|값2
                    for ci, f in enumerate(row_fields):
                        fid = f["field_id"]
                        val = str(filled.get(fid, ""))
                        label_cell = row.cells[ci * 2]
                        value_cell = row.cells[ci * 2 + 1]
                        p = label_cell.paragraphs[0]
                        run = p.add_run(f["label"])
                        run.bold = True
                        run.font.size = Pt(9)
                        _shade_cell(label_cell, "F8F8F8")
                        value_cell.paragraphs[0].add_run(val).font.size = Pt(10)

            # 테이블 후 간격
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)

        # 데이터 테이블 필드
        for f in table_fields:
            fid = f["field_id"]
            cols = f.get("table_columns", ["항목", "내용"])
            val = filled.get(fid, [])
            if isinstance(val, str):
                try:
                    val = json.loads(val)
                except:
                    val = []
            if not isinstance(val, list):
                val = []

            # 라벨
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(4)
            run = lp.add_run(f["label"])
            run.bold = True
            run.font.size = Pt(9)

            # 데이터 테이블
            dtbl = doc.add_table(rows=1, cols=len(cols))
            dtbl.style = 'Table Grid'
            dtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 헤더
            for ci, col in enumerate(cols):
                cell = dtbl.rows[0].cells[ci]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(col)
                run.bold = True
                run.font.size = Pt(9)
                _shade_cell(cell, "E8EDF2")
            # 데이터 행
            for row_data in val:
                dr = dtbl.add_row()
                for ci, col in enumerate(cols):
                    cell_val = ""
                    if isinstance(row_data, dict):
                        cell_val = str(row_data.get(col, ""))
                    elif isinstance(row_data, list) and ci < len(row_data):
                        cell_val = str(row_data[ci])
                    dr.cells[ci].paragraphs[0].add_run(cell_val).font.size = Pt(9)
                    dr.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not val:
                dr = dtbl.add_row()
                for ci in range(len(cols)):
                    dr.cells[ci].paragraphs[0].add_run("").font.size = Pt(9)

            # 데이터 테이블 후 간격
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)

    # closing 영역
    doc.add_paragraph()
    # 신청 문구 (있으면)
    if filled.get("application_date") or filled.get("applicant_name") or filled.get("recipient"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("본인은 관련 규정에 따라 위와 같이 신청합니다.")
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(filled.get("application_date", "　　년　　월　　일")).font.size = Pt(12)
        doc.add_paragraph()
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.add_run(f"신청인  {filled.get('applicant_name', '　　　　')}  (인)")
        if filled.get("recipient"):
            doc.add_paragraph()
            p4 = doc.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p4.add_run(filled.get("recipient", ""))
            run.bold = True
            run.font.size = Pt(14)

    doc.save(output_path)


def _shade_cell(cell, color: str):
    """DOCX 셀 배경색"""
    from docx.oxml.ns import qn
    shading = cell._element.get_or_add_tcPr().makeelement(
        qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    cell._element.get_or_add_tcPr().append(shading)


# ── Static ──
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
async def index():
    return FileResponse("static/index.html", headers={"Cache-Control": "no-cache, no-store, must-revalidate"})

@app.get("/download/templates")
async def download_templates():
    zip_path = Path(__file__).parent.parent / "templates_export.zip"
    if zip_path.exists():
        return FileResponse(str(zip_path), filename="templates_export.zip", media_type="application/zip")
    return JSONResponse({"error": "파일 없음"}, status_code=404)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
