import os
import io
import json
import zipfile
from pathlib import Path
from datetime import datetime

from fastapi import APIRouter, Query, HTTPException
from fastapi.responses import StreamingResponse, FileResponse, JSONResponse

from api.services.template_renderer import render_template

router = APIRouter()

DATA_DIR = Path(os.getenv("DATA_DIR", "./data"))


@router.get("/export/package")
async def export_package(nums: str = Query(default=None)):
    """전체 또는 선택 서식 ZIP 다운로드."""
    registry = _load_registry()

    if nums:
        selected = [n.strip() for n in nums.split(",")]
    else:
        selected = list(registry["forms"].keys())

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        export_registry = {"forms": {}}

        for num in selected:
            info = registry["forms"].get(num)
            if not info:
                continue

            form_id = info["form_id"]
            folder_name = f"{form_id}"
            schema_path = DATA_DIR / info["schema_path"]
            template_path = DATA_DIR / info["template_path"]

            if schema_path.exists():
                zf.write(schema_path, f"{folder_name}/schema.json")
            if template_path.exists():
                zf.write(template_path, f"{folder_name}/template.j2")

            export_registry["forms"][num] = info

        zf.writestr(
            "registry.json",
            json.dumps(export_registry, ensure_ascii=False, indent=2),
        )

        readme = f"""NUMP 서식 템플릿 패키지
생성일: {datetime.now().strftime('%Y-%m-%d')}
서식 수: {len(selected)}개
생성 도구: NUMP Template Converter

설치 방법:
1. 이 ZIP을 챗봇 서버의 data/ 디렉토리에 압축 해제
2. data/registry.json 위치 확인
3. 챗봇 서버 재시작

문의: support@nump.co.kr
"""
        zf.writestr("README.txt", readme)

    buf.seek(0)
    filename = f"templates_package_{datetime.now().strftime('%Y%m%d')}.zip"
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.post("/export/{num}/save")
async def save_single(num: int, body: dict = None):
    """개별 서식을 지정 포맷으로 서식 폴더에 저장.
    body: {"format": "pdf" | "docx" | "html"}
    """
    fmt = (body or {}).get("format", "pdf")
    if fmt not in ("pdf", "docx", "html"):
        raise HTTPException(status_code=400, detail="지원 포맷: pdf, docx, html")

    registry = _load_registry()
    info = registry["forms"].get(str(num))
    if not info:
        raise HTTPException(status_code=404, detail="서식을 찾을 수 없습니다")

    schema_path = DATA_DIR / info["schema_path"]
    template_path = DATA_DIR / info["template_path"]
    if not schema_path.exists() or not template_path.exists():
        raise HTTPException(status_code=404, detail="템플릿 파일이 없습니다")

    schema = json.loads(schema_path.read_text(encoding="utf-8"))
    template_str = template_path.read_text(encoding="utf-8")

    # 빈 서식 렌더링
    from api.services.template_renderer import _is_closing_field
    sample_data = {"form_name": schema.get("form_name", "")}
    for section in schema.get("sections", []):
        for field in section.get("fields", []):
            if _is_closing_field(field):
                continue
            fid = field["field_id"]
            if field.get("type") == "table":
                sample_data[fid] = []
            else:
                sample_data[fid] = ""

    html = render_template(template_str, sample_data)
    form_id = schema.get("form_id", info.get("form_id", "form"))
    form_dir = schema_path.parent

    if fmt == "html":
        out_path = form_dir / f"{form_id}.html"
        out_path.write_text(html, encoding="utf-8")

    elif fmt == "pdf":
        out_path = form_dir / f"{form_id}.pdf"
        try:
            _html_to_pdf(html, str(out_path))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF 생성 실패: {str(e)}")

    elif fmt == "docx":
        out_path = form_dir / f"{form_id}.docx"
        try:
            _schema_to_docx(schema, str(out_path))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX 생성 실패: {str(e)}")

    return JSONResponse({
        "status": "saved",
        "format": fmt,
        "path": str(out_path),
        "filename": out_path.name,
    })


@router.get("/export/{num}/download/{fmt}")
async def download_single(num: int, fmt: str):
    """저장된 개별 서식 파일 다운로드."""
    if fmt not in ("pdf", "docx", "html"):
        raise HTTPException(status_code=400, detail="지원 포맷: pdf, docx, html")

    registry = _load_registry()
    info = registry["forms"].get(str(num))
    if not info:
        raise HTTPException(status_code=404, detail="서식을 찾을 수 없습니다")

    schema_path = DATA_DIR / info["schema_path"]
    schema = json.loads(schema_path.read_text(encoding="utf-8"))
    form_id = schema.get("form_id", info.get("form_id", "form"))
    form_dir = schema_path.parent
    out_path = form_dir / f"{form_id}.{fmt}"

    if not out_path.exists():
        raise HTTPException(status_code=404, detail=f"{fmt.upper()} 파일이 없습니다. 먼저 저장하세요.")

    media_types = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "html": "text/html",
    }
    return FileResponse(
        str(out_path),
        media_type=media_types[fmt],
        filename=out_path.name,
    )


def _html_to_pdf(html: str, output_path: str):
    """pymupdf(fitz)로 HTML → PDF 변환."""
    import fitz
    # HTML을 임시 파일로 저장 후 PDF 생성
    story = fitz.Story(html=html)
    writer = fitz.DocumentWriter(output_path)
    # A4 크기
    mediabox = fitz.paper_rect("a4")
    where = mediabox + fitz.Rect(36, 36, -36, -36)  # 여백
    more = True
    while more:
        dev = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(dev)
        writer.end_page()
    writer.close()


def _schema_to_docx(schema: dict, output_path: str):
    """schema 기반으로 DOCX 서식 직접 생성. 테이블 포함."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from api.services.template_renderer import _is_closing_field

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(schema.get("form_name", "서식"))
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()  # 빈 줄

    for section in schema.get("sections", []):
        # 섹션 제목
        sp = doc.add_paragraph()
        run = sp.add_run(f"■ {section['section_name']}")
        run.bold = True
        run.font.size = Pt(11)

        # 테이블 필드 수집 (closing 제외)
        table_fields = [f for f in section.get("fields", []) if not _is_closing_field(f)]
        if not table_fields:
            continue

        # 테이블 생성
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for field in table_fields:
            ftype = field.get("type", "text")
            label = field.get("label", "")

            if ftype == "table":
                # 데이터 테이블 헤더
                cols = field.get("table_columns", ["항목", "내용"])
                row = tbl.add_row()
                row.cells[0].text = label
                row.cells[0].merge(row.cells[1])
                row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else None
                # 컬럼 헤더 행
                row2 = tbl.add_row()
                row2.cells[0].text = " | ".join(cols)
                row2.cells[0].merge(row2.cells[1])
                # 빈 데이터 행
                row3 = tbl.add_row()
                row3.cells[0].text = ""
                row3.cells[0].merge(row3.cells[1])
            else:
                row = tbl.add_row()
                # 라벨 셀
                cell_label = row.cells[0]
                cell_label.text = label
                for p in cell_label.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
                # 값 셀 (빈칸)
                cell_value = row.cells[1]
                if ftype == "textarea":
                    cell_value.text = "\n\n\n"  # 여러 줄 공간
                else:
                    cell_value.text = ""

        doc.add_paragraph()  # 섹션 간 간격

    # closing 영역
    doc.add_paragraph()
    closing_fields = []
    for section in schema.get("sections", []):
        for f in section.get("fields", []):
            if _is_closing_field(f):
                closing_fields.append(f)

    for f in closing_fields:
        label = f.get("label", "")
        fid = f.get("field_id", "")
        if "첨부" in label or "attach" in fid:
            doc.add_paragraph(f"※ {label}: ")
        elif "수신" in label or "귀하" in label or "recipient" in fid:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("○○○○ 귀하")
            run.bold = True
            run.font.size = Pt(14)

    # 신청 문구
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("본인은 관련 규정에 따라 위와 같이 신청합니다.")
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("　　년　　월　　일")
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.add_run("신청인 　　　　 (인)")

    doc.save(output_path)


def _load_registry() -> dict:
    path = DATA_DIR / "registry.json"
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {"forms": {}}
